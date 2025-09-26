import pdfplumber
import docx
import spacy
import os
import re

from datetime import datetime


def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join(para.text for para in doc.paragraphs)

# def load_resume_text(file_path):
#     if file_path.endswith(".pdf"):
#         return extract_text_from_pdf(file_path)
#     elif file_path.endswith(".docx"):
#         return extract_text_from_docx(file_path)
#     else:
#         raise ValueError("Unsupported file format")

def load_resume_text(file_input):
    text = ""

    # Case 1: string file path
    if isinstance(file_input, str):
        if file_input.endswith(".pdf"):
            with pdfplumber.open(file_input) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() + "\n"
        elif file_input.endswith(".docx"):
            doc = docx.Document(file_input)
            text = "\n".join([para.text for para in doc.paragraphs])

    # Case 2: Streamlit UploadedFile
    else:
        if file_input.name.endswith(".pdf"):
            with pdfplumber.open(file_input) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() + "\n"
        elif file_input.name.endswith(".docx"):
            doc = docx.Document(file_input)
            text = "\n".join([para.text for para in doc.paragraphs])

    return text
    
# def split_sections(text):
#     lines = text.split("\n")
#     sections = {}
#     current_section = None
#     buffer = []

#     for line in lines:
#         line = line.strip()

#         # Heuristics: Identify section headings
#         if re.match(r"^[A-Z][A-Za-z\s]{2,25}$", line) and line.isupper() or line.lower() in [
#             "experience", "work experience", "professional experience", "employment history"
#         ]:
#             if current_section:
#                 sections[current_section] = "\n".join(buffer).strip()
#                 buffer = []
#             current_section = line.lower()
#         elif current_section:
#             buffer.append(line)

#     if current_section:
#         sections[current_section] = "\n".join(buffer).strip()

#     return sections


# def extract_experience_section(text):
#     sections = split_sections(text)
#     for key in sections:
#         if "experience" in key or "employment" in key:
#             return sections[key]
#     return "Experience section not found."



# Define common resume section headings
SECTION_KEYWORDS = [
    "experience", "work experience", "professional experience", "employment history",
    "education", "skills", "projects", "certifications", "awards",
    "publications", "volunteer", "interests", "languages", "summary", "profile"
]

MONTH_NAMES = r'(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)'
DATE_PTN = rf'(?:\d{{1,2}}/\d{{4}}|{MONTH_NAMES}\s+\d{{4}})'
RANGE_PTN = re.compile(
    rf'(?P<from>{DATE_PTN})\s*(?:to|-)\s*(?P<to>(?:present|now|{DATE_PTN}))',
    re.IGNORECASE
)

def split_sections(text):
    lines = text.split("\n")
    sections = {}
    current_section = None
    buffer = []

    for line in lines:
        line_stripped = line.strip()

        # Identify section headings strictly
        if (
            re.match(r"^[A-Z][A-Z\s]{2,40}$", line_stripped)  # ALL CAPS headings
            or line_stripped.lower() in SECTION_KEYWORDS
        ):
            if current_section:
                sections[current_section] = "\n".join(buffer).strip()
                buffer = []
            current_section = line_stripped.lower()
        elif current_section:
            buffer.append(line_stripped)

    if current_section:
        sections[current_section] = "\n".join(buffer).strip()

    return sections


def extract_experience_section(text):
    sections = split_sections(text)

    # Return only the experience section
    for key, value in sections.items():
        if any(exp in key for exp in ["experience", "employment"]):
            return value.strip()

    return "Experience section not found."


def extract_education_section(text):
    sections = split_sections(text)

    # Return only the experience section
    for key, value in sections.items():
        if any(exp in key for exp in ["education", "academics"]):
            return value.strip()

    return "Experience section not found."

def extract_projects_section(text):
    sections = split_sections(text)

    # Return only the experience section
    for key, value in sections.items():
        if any(exp in key for exp in ["projects", "project"]):
            return value.strip()

    return "Project section not found."


def extract_structured_projects(raw_text):
    lines = [line.strip() for line in raw_text.split("\n") if line.strip()]
    projects = []
    i = 0
    while i < len(lines):
        # Treat line as project title if it's short (<120 chars) and has few periods
        if len(lines[i]) < 120 and lines[i].count('.') <= 1:
            project_name = lines[i]
            i += 1
            description_lines = []
            
            # Collect description lines until next potential title
            while i < len(lines) and not (len(lines[i]) < 120 and lines[i].count('.') <= 1):
                description_lines.append(lines[i])
                i += 1
            
            description = " ".join(description_lines).strip()
            projects.append({"project_name": project_name, "description": description})
        else:
            i += 1
    return projects


# def extract_structured_education(text: str):
#     # Pre-clean text to normalize spaces
#     text = re.sub(r'\s+', ' ', text)
    
#     # Split into chunks by major separators like "Core topics", "Degree Classification", "Graduated"
#     # This helps separate multiple education entries
#     entries = re.split(r'(?=University|Institute|College)', text, flags=re.IGNORECASE)
#     education_list = []

#     for entry in entries:
#         if not entry.strip():
#             continue

#         # Extract university/institute name (flexible patterns)
#         uni_match = re.search(
#             r'([A-Z][A-Za-z\s&,.]*?(University|Institute|College|School|Academy)[A-Za-z\s&,.]*)',
#             entry, re.IGNORECASE
#         )
#         university = uni_match.group(1).strip() if uni_match else None

#         # Extract degree (look for keywords like Master, Bachelor, PhD, MSc, BTech etc.)
#         degree_match = re.search(
#             r'(Master|Bachelor|Ph\.?D|MSc|BSc|BTech|MTech|MBA|MA|MS)\s*of?\s*[^-–|,]+',
#             entry, re.IGNORECASE
#         )
#         degree = degree_match.group(0).strip() if degree_match else None

#         # Extract start and end dates (handles multiple formats)
#         date_match = re.search(
#             r'([A-Za-z]{3,9}\s?\d{4})\s*[-–to]+\s*([A-Za-z]{3,9}\s?\d{4})',
#             entry, re.IGNORECASE
#         )
#         start_date, end_date = (date_match.group(1), date_match.group(2)) if date_match else (None, None)

#         # Extract CGPA/GPA if present
#         cgpa_match = re.search(r'(?:CGPA|GPA)\s*(?:of|:)?\s*([\d\.]+)', entry, re.IGNORECASE)
#         cgpa = cgpa_match.group(1) if cgpa_match else None

#         # Extract classification (like Distinction, First Class)
#         classification_match = re.search(
#             r'(?:Degree Classification|Classification|Grade)\s*[:\-]?\s*([A-Za-z\s]+)',
#             entry, re.IGNORECASE
#         )
#         classification = classification_match.group(1).strip() if classification_match else None

#         # Build structured result
#         education_list.append({
#             "university": university,
#             "degree": degree,
#             "start_date": start_date,
#             "end_date": end_date,
#             "cgpa": cgpa,
#             "classification": classification
#         })

#     return education_list

def extract_structured_education(text: str):
    # Normalize spaces
    text = re.sub(r'\s+', ' ', text)

    # Split into likely education entries
    entries = re.split(r'(?=University|Institute|College|School|Academy)', text, flags=re.IGNORECASE)
    education_list = []

    for entry in entries:
        entry = entry.strip()
        if not entry:
            continue

        # 1️⃣ University/Institution Name
        uni_match = re.search(
            r'([A-Z][A-Za-z&.,\-\s]*?(University|Institute|College|School|Academy)(?:\b|[A-Z]))',
            entry,
            re.IGNORECASE,
        )
        university = uni_match.group(1).strip() if uni_match else None

        # 2️⃣ Degree Level
        degree_level_match = re.search(
            r'\b(High School|Secondary School|Diploma|Associate|Bachelor(?:\'?s)?|Master(?:\'?s)?|Ph\.?D|Doctorate|MPhil|MBA|MSc|MS|BSc|BTech|MTech|MA|BA)\b',
            entry,
            re.IGNORECASE,
        )
        degree_level = degree_level_match.group(1).strip() if degree_level_match else None
        if degree_level:
            # Normalize to a cleaner form (remove "'s")
            degree_level = re.sub(r"\'s$", "", degree_level, flags=re.IGNORECASE).title()

        # 3️⃣ Subject/Field of Study
        # Pattern: (Bachelor/Master/etc.) [of|in] SUBJECT ...
        subject_match = re.search(
            r'(?:of|in)\s+([A-Z][A-Za-z&\-\s]{2,})',
            entry,
            re.IGNORECASE,
        )
        degree_subject = None
        if subject_match:
            raw = subject_match.group(1).strip()
            # Remove trailing noise like 'with', 'specialization', 'honours', etc.
            degree_subject = re.split(
                r'\b(with|specialization|honours|honors|distinction|major|minor|degree|program)\b',
                raw,
                flags=re.IGNORECASE
            )[0].strip()
            # Remove trailing commas or dashes
            degree_subject = re.sub(r'[,–-]+$', '', degree_subject).strip()

        # 4️⃣ Dates
        date_match = re.search(
            r'([A-Za-z]{3,9}\s?\d{4})\s*[-–to]+\s*([A-Za-z]{3,9}\s?\d{4})',
            entry,
            re.IGNORECASE,
        )
        start_date, end_date = (date_match.group(1), date_match.group(2)) if date_match else (None, None)

        # 5️⃣ GPA / CGPA
        cgpa_match = re.search(r'(?:CGPA|GPA)\s*(?:of|:)?\s*([\d\.]+)', entry, re.IGNORECASE)
        cgpa = cgpa_match.group(1) if cgpa_match else None

        # 6️⃣ Classification (Distinction, First Class, etc.)
        classification_match = re.search(
            r'(?:Degree Classification|Classification|Grade|Division)\s*[:\-]?\s*([A-Za-z\s]+)',
            entry,
            re.IGNORECASE,
        )
        classification = classification_match.group(1).strip() if classification_match else None

        # Append structured result
        education_list.append({
            "university": university,
            "degree_level": degree_level,
            "degree_subject": degree_subject,
            "start_date": start_date,
            "end_date": end_date,
            "cgpa": cgpa,
            "classification": classification,
        })

    return education_list


def normalize_date(date_str):
    if not date_str: 
        return None
    s = date_str.strip().lower()
    if s in ("present", "now"):
        return "Present"
    # Try common formats
    for fmt in ("%m/%Y", "%b %Y", "%B %Y"):
        try:
            return datetime.strptime(date_str.strip(), fmt).strftime("%Y-%m")
        except Exception:
            pass
    return date_str.strip()


def _find_prev_nonempty_line(lines, idx):
    for j in range(idx-1, -1, -1):
        if lines[j].strip():
            return lines[j].strip(), j
    return None, None




def extract_structured_experience_entries(text):
    entries = []
    if not text:
        return entries

    # split lines and prepare start positions for mapping character indices -> line index
    lines = text.splitlines()
    line_start_positions = []
    pos = 0
    for ln in lines:
        line_start_positions.append(pos)
        pos += len(ln) + 1  # +1 for the newline that was stripped by splitlines()

    # find all date-range matches
    matches = list(RANGE_PTN.finditer(text))
    if not matches:
        return entries

    # helper to map char index to line index
    def char_index_to_line_index(char_idx):
        # find rightmost line_start_positions[i] <= char_idx
        for i in range(len(line_start_positions)-1, -1, -1):
            if line_start_positions[i] <= char_idx:
                return i
        return 0

    # precompute line indices for each match
    match_line_indices = [char_index_to_line_index(m.start()) for m in matches]

    for idx, match in enumerate(matches):
        line_idx = match_line_indices[idx]
        line_text = lines[line_idx].rstrip()
        line_start_pos = line_start_positions[line_idx]
        rel_start_in_line = match.start() - line_start_pos

        # text before the date on the same line
        before_date = line_text[:rel_start_in_line].strip()

        # determine position and company
        position = ""
        company = ""

        # 1) Tech style: "Position | Company  Apr 2024 to Jan 2025"
        if '|' in before_date:
            left, right = [p.strip() for p in before_date.split('|', 1)]
            position = left
            company = right

        else:
            # find previous non-empty line (maybe position)
            prev_line, prev_idx = _find_prev_nonempty_line(lines, line_idx)

            if before_date:
                # company on same line; position likely on previous line
                company = before_date
                # use previous line as position if it doesn't look like a section heading or a date line
                if prev_line and not re.search(r'\d{4}', prev_line) and prev_line.lower() not in SECTION_KEYWORDS:
                    position = prev_line
            else:
                # before_date empty -> company might be previous line, and position one line above that
                if prev_line and prev_line.lower() not in SECTION_KEYWORDS:
                    company = prev_line
                    pos_prev_line, pos_prev_idx = _find_prev_nonempty_line(lines, prev_idx)
                    if pos_prev_line and pos_prev_line.lower() not in SECTION_KEYWORDS:
                        position = pos_prev_line

        # fallback: if still empty, try to infer from header_line segments split by multiple spaces
        if not company and before_date:
            parts = re.split(r'\s{2,}', before_date)
            if parts:
                company = parts[-1].strip()
                if len(parts) > 1:
                    position = parts[0].strip()

        # sanitize: if any trailing date fragments ended up in company, strip them
        company = re.sub(r'\s{2,}$', '', company).strip()
        position = position.strip()

        # get summary: from next line after this line up to next match's line (exclusive)
        next_line_idx = match_line_indices[idx+1] if idx+1 < len(match_line_indices) else len(lines)
        summary_lines = lines[line_idx+1: next_line_idx]
        # drop any following section headings that might appear at start of summary
        cleaned_summary = []
        for ln in summary_lines:
            s = ln.strip()
            if not s:
                continue
            # stop if a known section heading starts the summary chunk
            if s.lower() in SECTION_KEYWORDS or re.match(r'^[A-Z][A-Z\s]{2,40}$', s):
                break
            cleaned_summary.append(s)
        summary = " ".join(cleaned_summary).strip()

        entry = {
            "company_name": company or "",
            "position": position or "",
            "from": normalize_date(match.group("from")),
            "to": normalize_date(match.group("to")),
            "summary": summary
        }
        entries.append(entry)

    return entries